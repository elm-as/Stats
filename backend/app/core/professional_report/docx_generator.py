"""
Génération DOCX avec python-docx — style mémoire universitaire.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from app.core.professional_report.builder import ReportContent, ReportSection


ACCENT_RGB = RGBColor(0x06, 0xB6, 0xD4)
SECONDARY_RGB = RGBColor(0x1E, 0x3A, 0x8A)
GRAY_TXT = RGBColor(0x1F, 0x29, 0x37)
GRAY_LIGHT = RGBColor(0x6B, 0x72, 0x80)

SEVERITY_RGB = {
    "critical": RGBColor(0xDC, 0x26, 0x26),
    "warning": RGBColor(0xD9, 0x77, 0x06),
    "success": RGBColor(0x05, 0x96, 0x69),
    "methodological": RGBColor(0x7C, 0x3A, 0xED),
    "info": RGBColor(0x25, 0x63, 0xEB),
}


def _strip_md(text: str) -> tuple[str, list[tuple[int, int, str]]]:
    """Extrait les emphases : retourne (texte_clean, [(start, end, style)])."""
    if not text:
        return "", []

    runs = []
    out = []
    cursor = 0

    pattern = re.compile(r"(\*\*([^*]+)\*\*|`([^`]+)`)")
    last = 0
    for m in pattern.finditer(text):
        out.append(text[last:m.start()])
        cursor += m.start() - last

        if m.group(2) is not None:  # **bold**
            out.append(m.group(2))
            runs.append((cursor, cursor + len(m.group(2)), "bold"))
            cursor += len(m.group(2))
        elif m.group(3) is not None:  # `code`
            out.append(m.group(3))
            runs.append((cursor, cursor + len(m.group(3)), "code"))
            cursor += len(m.group(3))

        last = m.end()

    out.append(text[last:])
    return "".join(out), runs


def _add_styled(paragraph, text: str, base_size: int = 11, color: RGBColor | None = None):
    clean, runs = _strip_md(text)
    if not runs:
        run = paragraph.add_run(clean)
        run.font.size = Pt(base_size)
        if color:
            run.font.color.rgb = color
        return

    cursor = 0
    for start, end, style in runs:
        # Texte avant
        if start > cursor:
            r = paragraph.add_run(clean[cursor:start])
            r.font.size = Pt(base_size)
            if color:
                r.font.color.rgb = color
        # Texte stylisé
        r = paragraph.add_run(clean[start:end])
        r.font.size = Pt(base_size)
        if style == "bold":
            r.bold = True
            if color:
                r.font.color.rgb = color
        elif style == "code":
            r.font.name = "Consolas"
            r.font.color.rgb = ACCENT_RGB
        else:
            if color:
                r.font.color.rgb = color
        cursor = end

    if cursor < len(clean):
        r = paragraph.add_run(clean[cursor:])
        r.font.size = Pt(base_size)
        if color:
            r.font.color.rgb = color


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    if level == 0:
        run.font.size = Pt(28)
        run.font.color.rgb = SECONDARY_RGB
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = ACCENT_RGB
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY_TXT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY_TXT
    return p


def _add_section(doc, section: ReportSection, level: int = 1):
    _add_heading(doc, section.title, level=level)

    if section.body:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_styled(p, section.body, base_size=11, color=GRAY_TXT)

    if section.bullets:
        for b in section.bullets:
            p = doc.add_paragraph(style="List Bullet")
            _add_styled(p, b, base_size=10.5, color=GRAY_TXT)

    # Insights
    for ins in section.insights or []:
        sev = ins.get("severity", "info")
        col = SEVERITY_RGB.get(sev, SEVERITY_RGB["info"])

        p = doc.add_paragraph()
        run = p.add_run(f"▸ {ins.get('title', '')}")
        run.bold = True
        run.font.color.rgb = col
        run.font.size = Pt(11)

        if ins.get("message"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.6)
            _add_styled(p2, ins["message"], base_size=10, color=GRAY_TXT)

        if ins.get("suggestion"):
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(0.6)
            r = p3.add_run("→ Recommandation : ")
            r.italic = True
            r.font.color.rgb = ACCENT_RGB
            r.font.size = Pt(9.5)
            _add_styled(p3, ins["suggestion"], base_size=9.5, color=GRAY_LIGHT)

    # Table
    if section.table:
        headers = section.table.get("headers", [])
        rows = section.table.get("rows", [])
        if rows and headers:
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = "Light Grid Accent 1"
            # Headers
            hdr_cells = t.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = ""
                p = hdr_cells[i].paragraphs[0]
                r = p.add_run(str(h))
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Rows
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    cell = t.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    r = p.add_run(str(val))
                    r.font.size = Pt(9.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for sub in section.subsections or []:
        _add_section(doc, sub, level=level + 1)


def generate_docx(content: ReportContent) -> bytes:
    """Génère un DOCX depuis un ReportContent. Retourne les bytes."""
    doc = Document()

    # Style global
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    _add_heading(doc, content.title, level=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(content.subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_RGB

    for _ in range(4):
        doc.add_paragraph()

    meta_tbl = doc.add_table(rows=0, cols=2)
    for k, v in [("Date", content.date), ("Auteur", content.author),
                 *((str(k), str(v)) for k, v in (content.metadata or {}).items())]:
        row = meta_tbl.add_row().cells
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(k)
        r.font.color.rgb = GRAY_LIGHT
        r.font.size = Pt(10)
        row[1].text = ""
        r2 = row[1].paragraphs[0].add_run(v)
        r2.bold = True
        r2.font.size = Pt(10)

    doc.add_page_break()

    # Executive summary
    if content.executive_summary:
        _add_heading(doc, "Résumé exécutif", level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_styled(p, content.executive_summary, base_size=11, color=GRAY_TXT)

    # Key findings
    if content.key_findings:
        _add_heading(doc, "Faits saillants", level=1)
        for kf in content.key_findings:
            p = doc.add_paragraph(style="List Bullet")
            _add_styled(p, kf, base_size=10.5, color=GRAY_TXT)

    # Sections
    for sec in content.sections:
        _add_section(doc, sec, level=1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
